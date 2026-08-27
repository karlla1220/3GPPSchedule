import unittest
from pathlib import Path
from zipfile import ZipFile
from xml.etree.ElementTree import Element, SubElement
from unittest.mock import MagicMock

from docx import Document

from models import RoomInfo
from parser import (
    _determine_time_block_index,
    _get_cell_text,
    extract_meeting_location,
    find_chair_notes_docx,
    parse_docx,
)
from session_parser import _extract_agenda_item_from_name, _slot_result_to_sessions

_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _write_word_package(path: Path, paragraphs: list[str]) -> None:
    body = "".join(
        f"<w:p><w:r><w:t>{text}</w:t></w:r></w:p>" for text in paragraphs
    )
    with ZipFile(path, "w") as package:
        package.writestr(
            "word/document.xml",
            f'''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="{_NS}"><w:body>{body}</w:body></w:document>''',
        )


def test_extract_meeting_location_reads_docx_and_docm_through_ooxml(tmp_path):
    for suffix in (".docx", ".docm"):
        path = tmp_path / f"Chair notes RAN1#126_v00{suffix}"
        _write_word_package(
            path,
            [
                "3GPP TSG RAN WG1 #126",
                "Maastricht, NL, Aug 24th-28th, 2026",
            ],
        )

        assert (
            extract_meeting_location(path)
            == "Maastricht, NL, Aug 24th-28th, 2026"
        )


def test_find_chair_notes_docx_accepts_docm_for_current_meeting_only(tmp_path):
    old = tmp_path / "Chair notes RAN1#125_v12.docx"
    current = tmp_path / "Chair notes RAN1#126_v00.docm"
    _write_word_package(old, ["Shanghai, CN, May 18th-22nd, 2026"])
    _write_word_package(current, ["Maastricht, NL, Aug 24th-28th, 2026"])

    assert find_chair_notes_docx(tmp_path, meeting_id="ran1#126") == current


def test_extract_meeting_location_returns_none_for_malformed_docm(tmp_path):
    path = tmp_path / "Chair notes RAN1#126_v00.docm"
    path.write_bytes(b"not-an-ooxml-package")

    assert extract_meeting_location(path) is None


def _make_cell_xml(*paragraphs):
    """Build a mock cell whose _tc contains paragraph/run XML.

    Each paragraph is a list of (text, strike) tuples where *strike*
    is ``True`` when the run should carry ``<w:strike/>``.
    """
    tc = Element(f"{{{_NS}}}tc")
    for runs in paragraphs:
        p = SubElement(tc, f"{{{_NS}}}p")
        for text, strike in runs:
            r = SubElement(p, f"{{{_NS}}}r")
            if strike:
                rpr = SubElement(r, f"{{{_NS}}}rPr")
                SubElement(rpr, f"{{{_NS}}}strike")
            t = SubElement(r, f"{{{_NS}}}t")
            t.text = text
    cell = MagicMock()
    cell._tc = tc
    return cell


class GetCellTextTests(unittest.TestCase):
    def test_normal_text_preserved(self):
        cell = _make_cell_xml([("hello", False), (" world", False)])
        self.assertEqual(_get_cell_text(cell), "hello world")

    def test_strike_text_marked(self):
        cell = _make_cell_xml([
            ("Xiaodong (150) 6GR ", False),
            (".10.5.1.2 (30)", True),
            (" .10.5.1.3 (60)", False),
        ])
        self.assertEqual(
            _get_cell_text(cell),
            "Xiaodong (150) 6GR ~~.10.5.1.2 (30)~~ .10.5.1.3 (60)",
        )

    def test_all_strike_marked(self):
        cell = _make_cell_xml([("deleted", True)])
        self.assertEqual(_get_cell_text(cell), "~~deleted~~")

    def test_multi_paragraph(self):
        cell = _make_cell_xml(
            [("line1", False)],
            [("line2", False)],
        )
        self.assertEqual(_get_cell_text(cell), "line1\nline2")

    def test_dstrike_marked(self):
        """Double-strikethrough (dstrike) should also be marked."""
        tc = Element(f"{{{_NS}}}tc")
        p = SubElement(tc, f"{{{_NS}}}p")
        r = SubElement(p, f"{{{_NS}}}r")
        rpr = SubElement(r, f"{{{_NS}}}rPr")
        SubElement(rpr, f"{{{_NS}}}dstrike")
        t = SubElement(r, f"{{{_NS}}}t")
        t.text = "double-strike"
        cell = MagicMock()
        cell._tc = tc
        self.assertEqual(_get_cell_text(cell), "~~double-strike~~")

    def test_adjacent_strike_and_normal(self):
        """Strike markers should not bleed into adjacent normal text."""
        cell = _make_cell_xml([
            ("keep", False),
            ("remove", True),
            ("also keep", False),
        ])
        self.assertEqual(
            _get_cell_text(cell),
            "keep~~remove~~also keep",
        )


class DetermineTimeBlockIndexTests(unittest.TestCase):
    def test_matches_standard_block_start(self):
        self.assertEqual(
            _determine_time_block_index("08:30\n~\n10:30\n\n(120 min)"),
            0,
        )

    def test_matches_later_start_within_morning_block(self):
        self.assertEqual(
            _determine_time_block_index("09:00\n~\n10:30\n\n(90 min)"),
            0,
        )

    def test_does_not_match_break_start(self):
        self.assertIsNone(
            _determine_time_block_index("10:30\n~\n11:00\n\n(30 min)"),
        )


def test_parse_docx_applies_grid_span_on_vmerge_time_continuation(tmp_path):
    path = tmp_path / "vmerge-gridspan.docx"
    document = Document()
    table = document.add_table(rows=5, cols=4)

    table.cell(0, 1).merge(table.cell(0, 3)).text = "Thursday"

    table.cell(1, 0).text = "08:30 ~ 10:30 (120 min)"
    table.cell(1, 1).text = "Room A morning"
    table.cell(1, 2).text = "Room B morning"
    table.cell(1, 3).text = "Room C morning"

    table.cell(2, 0).merge(table.cell(2, 3)).text = "Lunch break: 13:00 ~ 14:30"

    table.cell(3, 0).text = "17:00 ~ 19:30 (150 min)"
    table.cell(3, 1).text = "Room A evening"
    table.cell(3, 2).text = "Room B evening"
    table.cell(3, 3).text = "Room C evening"
    table.cell(3, 0).merge(table.cell(4, 0))
    table.cell(4, 1).merge(table.cell(4, 3)).text = "Early dinner (60)"

    document.save(path)

    cells, _ = parse_docx(path)
    dinners = [cell for cell in cells if cell.text == "Early dinner (60)"]

    assert len(dinners) == 1
    assert dinners[0].time_block_index == 3
    assert dinners[0].room_indices == [0, 1, 2]
    assert dinners[0].fallback_start_time == "18:30"


def test_parse_docx_does_not_inherit_time_for_plain_blank_row(tmp_path):
    path = tmp_path / "blank-time-gridspan.docx"
    document = Document()
    table = document.add_table(rows=5, cols=4)

    table.cell(0, 1).merge(table.cell(0, 3)).text = "Thursday"

    table.cell(1, 0).text = "08:30 ~ 10:30 (120 min)"
    table.cell(1, 1).text = "Room A morning"
    table.cell(1, 2).text = "Room B morning"
    table.cell(1, 3).text = "Room C morning"

    table.cell(2, 0).merge(table.cell(2, 3)).text = "Lunch break: 13:00 ~ 14:30"

    table.cell(3, 0).text = "17:00 ~ 19:30 (150 min)"
    table.cell(3, 1).text = "Room A evening"
    table.cell(3, 2).text = "Room B evening"
    table.cell(3, 3).text = "Room C evening"
    table.cell(4, 1).merge(table.cell(4, 3)).text = "Unrelated note"

    document.save(path)

    cells, _ = parse_docx(path)

    assert all(cell.text != "Unrelated note" for cell in cells)


class AgendaExtractionTests(unittest.TestCase):
    def test_extracts_prefixed_agenda_item(self):
        name, agenda_item = _extract_agenda_item_from_name("9.6 R20 NTN-NR")
        self.assertEqual(name, "R20 NTN-NR")
        self.assertEqual(agenda_item, "9.6")

    def test_preserves_slash_shorthand_in_agenda_item(self):
        name, agenda_item = _extract_agenda_item_from_name("10.5.1.2/3")
        self.assertEqual(name, "10.5.1.2/3")
        self.assertEqual(agenda_item, "10.5.1.2/3")

    def test_slot_result_keeps_full_name_for_slash_shorthand(self):
        slot = MagicMock(
            day="Friday",
            time_block_start="08:30",
            time_block_end="10:30",
        )
        parsed = {
            "sessions": [
                {
                    "room_name": "RAN1_off1",
                    "name": "10.5.1.2/3",
                    "duration_minutes": 120,
                    "specified_start_time": None,
                    "chair": "Xiaodong",
                    "group_header": "6GR",
                    "agenda_item": None,
                }
            ]
        }
        day_rooms_map = {
            "Friday": [
                RoomInfo(name="Some online room", table_index=0, room_index_in_table=0),
                RoomInfo(name="Offline Room 1", table_index=1, room_index_in_table=0),
            ]
        }
        alias_to_name = {"RAN1_off1": "Offline Room 1"}

        sessions = _slot_result_to_sessions(parsed, slot, day_rooms_map, alias_to_name)

        self.assertEqual(len(sessions), 1)
        self.assertEqual(sessions[0].name, "10.5.1.2/3")
        self.assertEqual(sessions[0].agenda_item, "10.5.1.2/3")
        self.assertEqual(sessions[0].chair, "Xiaodong")
        self.assertEqual(sessions[0].group_header, "6GR")

    def test_merged_session_starts_after_latest_member_room(self):
        slot = MagicMock(
            day="Thursday",
            time_block_start="17:00",
            time_block_end="19:30",
        )
        parsed = {
            "sessions": [
                {
                    "room_name": "RAN1_main",
                    "name": "Main session",
                    "duration_minutes": 60,
                    "specified_start_time": None,
                },
                {
                    "room_name": "RAN1_brk1",
                    "name": "Breakout session",
                    "duration_minutes": 90,
                    "specified_start_time": None,
                },
                {
                    "room_name": "ALL_ONLINE",
                    "name": "Early dinner",
                    "duration_minutes": 60,
                    "specified_start_time": None,
                },
            ]
        }
        day_rooms_map = {
            "Thursday": [
                RoomInfo(name="Main room", table_index=0, room_index_in_table=0),
                RoomInfo(name="Breakout room", table_index=0, room_index_in_table=1),
            ]
        }
        alias_to_name = {
            "RAN1_main": "Main room",
            "RAN1_brk1": "Breakout room",
            "ALL_ONLINE": "Main room + Breakout room",
        }

        sessions = _slot_result_to_sessions(
            parsed, slot, day_rooms_map, alias_to_name
        )

        dinner = next(s for s in sessions if s.name == "Early dinner")
        self.assertEqual((dinner.start_time, dinner.end_time), ("18:30", "19:30"))

    def test_merged_session_uses_fallback_when_latest_room_would_overrun(self):
        slot = MagicMock(
            day="Thursday",
            time_block_start="17:00",
            time_block_end="19:30",
        )
        slot.sources = [
            MagicMock(
                entries=[
                    MagicMock(
                        room_label="Main room + Breakout room",
                        cell_text="Early dinner (60)",
                        fallback_start_time="18:30",
                    )
                ]
            )
        ]
        parsed = {
            "sessions": [
                {
                    "room_name": "RAN1_main",
                    "name": "Overlong room schedule",
                    "duration_minutes": 150,
                    "specified_start_time": None,
                },
                {
                    "room_name": "RAN1_brk1",
                    "name": "Shorter room schedule",
                    "duration_minutes": 90,
                    "specified_start_time": None,
                },
                {
                    "room_name": "ALL_ONLINE",
                    "name": "Early dinner",
                    "duration_minutes": 60,
                    "specified_start_time": None,
                },
            ]
        }
        day_rooms_map = {
            "Thursday": [
                RoomInfo(name="Main room", table_index=0, room_index_in_table=0),
                RoomInfo(name="Breakout room", table_index=0, room_index_in_table=1),
            ]
        }
        alias_to_name = {
            "RAN1_main": "Main room",
            "RAN1_brk1": "Breakout room",
            "ALL_ONLINE": "Main room + Breakout room",
        }

        sessions = _slot_result_to_sessions(
            parsed, slot, day_rooms_map, alias_to_name
        )

        dinner = next(s for s in sessions if s.name == "Early dinner")
        self.assertEqual((dinner.start_time, dinner.end_time), ("18:30", "19:30"))


if __name__ == "__main__":
    unittest.main()
