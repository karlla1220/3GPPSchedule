"""Parse DOCX tables to extract raw schedule cell data."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from models import TIME_BLOCKS, CellData, DAY_ORDER, RoomInfo


def extract_meeting_location(docx_path: Path) -> str | None:
    """Extract the meeting location line from a Chair notes DOCX.

    Looks for a line like "Gothenburg, SE, Feb. 9th ~ 13th, 2026"
    in the first few paragraphs of the document.

    Returns the location string if found, None otherwise.
    """
    doc = Document(str(docx_path))
    # Check first 10 paragraphs for a location-like line
    for para in doc.paragraphs[:10]:
        text = para.text.strip()
        if not text:
            continue
        # Location lines typically look like:
        #   "Dallas, USA, Nov 17th – 21st, 2025"
        #   "Athens, Greece, February 17th – 21st, 2025"
        # Allow both 2-letter and full country names.
        if re.search(
            r"^[A-Z][A-Za-z .'-]+,\s*[A-Z][A-Za-z .'-]+,\s*[A-Za-z]+\.?\s+\d",
            text,
        ):
            return text
    return None


def find_chair_notes_docx(dest_dir: Path = Path("downloads/Chair_notes")) -> Path | None:
    """Find the latest Chair notes document in the local directory.

    Looks for files with 'chair note' (case-insensitive) in the name,
    supporting .docx, .pptx, and .pdf extensions.
    Returns the one with the highest modification time.
    """
    supported_extensions = (".docx", ".pptx", ".pdf")
    chair_files = [
        f
        for ext in supported_extensions
        for f in dest_dir.glob(f"*{ext}")
        if "chair note" in f.name.lower() or "chair_note" in f.name.lower()
    ]
    if not chair_files:
        return None
    return max(chair_files, key=lambda f: f.stat().st_mtime)


# Namespace for OpenXML
_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_V_NS = "urn:schemas-microsoft-com:vml"
_WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"


def _get_grid_span(cell) -> int:
    """Get the horizontal grid span of a cell."""
    tc = cell._tc
    grid_span = tc.find(f".//{{{_NS}}}gridSpan")
    if grid_span is not None:
        return int(grid_span.get(f"{{{_NS}}}val"))
    return 1


def _dedupe_row_cells(row) -> list[tuple]:
    """Get deduplicated cells from a row with their grid positions.

    Returns list of (cell_text, col_start, col_end) tuples.
    col_start is 0-indexed, col_end is exclusive.

    python-docx's row.cells returns one entry per grid column position,
    repeating the same cell object for merged columns. We use enumerate
    to track the actual column index.

    Cells that are vertical-merge continuations (vMerge without
    val="restart") are emitted with empty text so that column positions
    remain correct while the content is not double-counted.
    """
    # Build a set of grid-column ranges that are vMerge continuations
    # by inspecting the raw XML <w:tc> elements in this row's <w:tr>.
    # python-docx's row.cells resolves merged cells to the original,
    # so we cannot compare tc element identity; we use column positions.
    tr = row._tr
    continuation_cols: set[int] = set()
    col_cursor = 0
    for tc_el in tr.findall(f"{{{_NS}}}tc"):
        tc_pr = tc_el.find(f"{{{_NS}}}tcPr")
        gs = 1
        is_continuation = False
        if tc_pr is not None:
            gs_el = tc_pr.find(f"{{{_NS}}}gridSpan")
            if gs_el is not None:
                gs = int(gs_el.get(f"{{{_NS}}}val"))
            vm_el = tc_pr.find(f"{{{_NS}}}vMerge")
            if vm_el is not None:
                val = vm_el.get(f"{{{_NS}}}val")
                if val != "restart":
                    is_continuation = True
        if is_continuation:
            for i in range(col_cursor, col_cursor + gs):
                continuation_cols.add(i)
        col_cursor += gs

    result = []
    row_seen = set()

    for col_pos, cell in enumerate(row.cells):
        tc_id = id(cell._tc)
        if tc_id in row_seen:
            continue
        row_seen.add(tc_id)

        span = _get_grid_span(cell)

        # Emit empty text for vMerge continuation cells
        if col_pos in continuation_cols:
            result.append(("", col_pos, col_pos + span))
        else:
            text = cell.text.strip()
            result.append((text, col_pos, col_pos + span))

    return result


def _is_break_row(cells: list[tuple]) -> bool:
    """Check if this row is a break (coffee/lunch).

    Real break rows span the entire table (1-2 cells).  Data rows may
    mention 'break' inside schedule content but have many more cells.
    """
    if not cells or len(cells) > 3:
        return False
    full_text = " ".join(c[0] for c in cells).lower()
    return "break" in full_text or "coffee" in full_text or "lunch" in full_text


def _is_footer_row(cells: list[tuple]) -> bool:
    """Check if this row is the end-of-day footer."""
    if not cells or len(cells) > 3:
        return False
    full_text = " ".join(c[0] for c in cells).lower()
    return "all sessions end" in full_text or "no exceptions" in full_text


def _is_metadata_row(cells: list[tuple]) -> bool:
    """Check if this row contains room metadata."""
    if not cells or len(cells) > 3:
        return False
    full_text = " ".join(c[0] for c in cells).lower()
    return "meeting rooms" in full_text


def _parse_day_header(cells: list[tuple]) -> dict[str, tuple[int, int]]:
    """Parse the header row to map day names to column ranges.

    Returns dict: day_name -> (col_start, col_end) inclusive/exclusive.
    """
    day_map = {}
    for text, col_start, col_end in cells:
        # Normalize day name
        for day in DAY_ORDER:
            if day.lower() in text.lower():
                if day not in day_map:
                    day_map[day] = (col_start, col_end)
                break
    return day_map


def _parse_room_code(line: str) -> str:
    """Extract a short room code from a line like 'RAN1_Off#1 (J1)' → 'J1'.

    Tries the last parenthesised token first (e.g. '(J1)').  If none found,
    falls back to the whole line.
    """
    # Find all parenthesised groups
    matches = re.findall(r"\(([^)]+)\)", line)
    if matches:
        # Use the last match – it's usually the room code
        # e.g. "Main session (F1/2/3, Level 2)" → "F1/2/3, Level 2" → take part before comma
        code = matches[-1].split(",")[0].strip()
        return code
    return line.strip()


def _normalize_color(color: str | None) -> str | None:
    """Normalize a hex color string to uppercase 6-digit form.

    Handles '#RRGGBB', 'RRGGBB', and near-matches like D9D9D9 vs D8D8D8.
    Returns None for missing / 'auto' values.
    """
    if not color or color.lower() == 'auto':
        return None
    color = color.lstrip('#').upper()
    if len(color) == 6:
        return color
    return None


def _colors_match(c1: str, c2: str, tolerance: int = 8) -> bool:
    """Check if two hex colors are close enough to be considered the same.

    Allows small differences (e.g. D9D9D9 vs D8D8D8) caused by theme
    resolution differences between VML and OOXML.
    """
    if c1 == c2:
        return True
    try:
        r1, g1, b1 = int(c1[0:2], 16), int(c1[2:4], 16), int(c1[4:6], 16)
        r2, g2, b2 = int(c2[0:2], 16), int(c2[2:4], 16), int(c2[4:6], 16)
        return (
            abs(r1 - r2) <= tolerance
            and abs(g1 - g2) <= tolerance
            and abs(b1 - b2) <= tolerance
        )
    except (ValueError, IndexError):
        return False


def _extract_textbox_rooms(doc: Document) -> list[dict]:
    """Extract room labels from TextBox shapes in the document.

    Each TextBox that sits above a table acts as a room column header.
    Returns list of dicts: {'name': str, 'color': str | None}.
    Color is the fill/background colour of the shape (normalised hex).
    """
    body = doc.element.body
    alt_contents = body.findall(f'.//{{{_MC_NS}}}AlternateContent')

    results: list[dict] = []
    for ac in alt_contents:
        # Collect text from <w:t> inside the AlternateContent
        texts = []
        from lxml import etree  # noqa: local import
        for t_el in ac.iter(f'{{{_NS}}}t'):
            if t_el.text:
                texts.append(t_el.text)
        raw = ''.join(texts).strip()
        if not raw:
            continue

        # De-duplicate doubled text ('J1J1' → 'J1')
        # The same text appears in both mc:Choice and mc:Fallback.
        if len(raw) % 2 == 0:
            half = len(raw) // 2
            if raw[:half] == raw[half:]:
                raw = raw[:half]

        # --- Determine fill colour ---
        fill_color: str | None = None

        # 1) Try VML fillcolor attribute (most reliable resolved value)
        for shape in ac.iter(f'{{{_V_NS}}}shape'):
            fc = shape.get('fillcolor')
            if fc:
                # VML may include " [id]" suffix, e.g. '#ffd966 [1943]'
                fc = fc.split()[0]
                fill_color = _normalize_color(fc)
                break

        # 2) Fallback: DrawingML a:solidFill > a:srgbClr (skip font colors)
        if fill_color is None:
            for sp_pr in ac.iter(f'{{{_WPS_NS}}}spPr'):
                for solid in sp_pr.iter(f'{{{_A_NS}}}solidFill'):
                    srgb = solid.find(f'{{{_A_NS}}}srgbClr')
                    if srgb is not None:
                        fill_color = _normalize_color(srgb.get('val'))
                        break
                if fill_color is not None:
                    break

        results.append({'name': raw, 'color': fill_color})

    return results


def _get_table_column_colors(
    table, day_columns: dict[str, tuple[int, int]],
    actual_rooms: dict[str, int],
) -> list[str]:
    """Collect the ordered set of distinct room colours used in a table.

    Scans data rows looking for days where the number of distinct cells
    equals the actual room count.  From those cells, extracts the fill
    colours in column order.  This avoids picking up colours from merged
    cells or empty-day cells.

    Returns an ordered list of unique colour hex strings, one per room.
    """
    valid_ranges = list(day_columns.items())

    for row in table.rows[1:]:
        tr = row._tr

        # Build (col_start, col_end, fill) for every tc in this row
        tc_info: list[tuple[int, int, str | None]] = []
        col_cursor = 0
        for tc_el in tr.findall(f'{{{_NS}}}tc'):
            tc_pr = tc_el.find(f'{{{_NS}}}tcPr')
            gs = 1
            fill = None
            if tc_pr is not None:
                gs_el = tc_pr.find(f'{{{_NS}}}gridSpan')
                if gs_el is not None:
                    gs = int(gs_el.get(f'{{{_NS}}}val'))
                shd = tc_pr.find(f'{{{_NS}}}shd')
                if shd is not None:
                    fill = _normalize_color(shd.get(f'{{{_NS}}}fill'))
            tc_info.append((col_cursor, col_cursor + gs, fill))
            col_cursor += gs

        # For each day, collect room-cell colours if cell count matches
        for day_name, (day_cs, day_ce) in valid_ranges:
            expected = actual_rooms.get(day_name, 0)
            if expected < 1:
                continue

            day_fills: list[str] = []
            for cs, ce, fill in tc_info:
                if cs >= day_cs and ce <= day_ce and cs > 0:
                    day_fills.append(fill or 'FFFFFF')

            if len(day_fills) == expected:
                # De-dup while preserving order
                seen: set[str] = set()
                ordered: list[str] = []
                for f in day_fills:
                    canon = None
                    for s in seen:
                        if _colors_match(f, s):
                            canon = s
                            break
                    if canon is None:
                        seen.add(f)
                        ordered.append(f)
                if len(ordered) == expected:
                    return ordered

    return []


def _match_rooms_to_table(
    textbox_rooms: list[dict],
    table,
    day_columns: dict[str, tuple[int, int]],
    actual_rooms: dict[str, int],
) -> list[str] | None:
    """Match TextBox room labels to a table using background colour matching.

    Returns an ordered list of room names for this table's columns,
    or None if matching fails.
    """
    if not textbox_rooms:
        return None

    unique_table_colors = _get_table_column_colors(
        table, day_columns, actual_rooms,
    )
    if not unique_table_colors:
        return None

    # For each unique table colour, find matching TextBox
    matched_names: list[str] = []
    for tc in unique_table_colors:
        for tb in textbox_rooms:
            if tb['color'] and _colors_match(tc, tb['color']):
                matched_names.append(tb['name'])
                break

    # Validate: matched count should equal the actual room count
    sample_day = next(iter(actual_rooms))
    num_rooms = actual_rooms[sample_day]
    if len(matched_names) == num_rooms:
        return matched_names

    return None


def _extract_room_names_from_doc(
    doc: Document,
) -> tuple[list[str] | None, list[str] | None]:
    """Extract online and offline room names from document metadata rows.

    Fallback used when TextBox colour matching is not available.

    Returns (online_rooms, offline_rooms) – either may be None.
    """
    for table in doc.tables:
        for row in table.rows:
            cells = _dedupe_row_cells(row)
            for text, _, _ in cells:
                if "meeting rooms" not in text.lower():
                    continue

                # Split into blocks on every "Meeting Rooms" header
                blocks: list[list[str]] = []
                current: list[str] = []
                for line in text.split("\n"):
                    stripped = line.strip()
                    if not stripped:
                        continue
                    if "meeting rooms" in stripped.lower():
                        if current:
                            blocks.append(current)
                        current = []
                    else:
                        current.append(stripped)
                if current:
                    blocks.append(current)

                online_rooms: list[str] | None = None
                offline_rooms: list[str] | None = None

                for block in blocks:
                    # Determine if this block is for offline rooms
                    is_offline = any("off" in ln.lower() for ln in block)
                    names = [_parse_room_code(ln) for ln in block]
                    if is_offline:
                        offline_rooms = names
                    else:
                        online_rooms = names

                return online_rooms, offline_rooms

    return None, None


def _get_table_preceding_paragraphs(doc: Document) -> dict[int, str]:
    """Map each table index (in doc.tables order) to preceding paragraph text.

    Walks through the document body elements in order, collecting paragraph
    text until a table is encountered.  The collected text is associated
    with that table's index.

    Returns:
        dict mapping table index (0-based, matching doc.tables) to the
        concatenated text of all paragraphs between the previous table
        (or document start) and this table.
    """
    body = doc.element.body
    table_contexts: dict[int, str] = {}
    current_paragraphs: list[str] = []
    table_idx = 0

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            # Extract text from paragraph XML element
            texts = []
            for t_el in child.iter(f'{{{_NS}}}t'):
                if t_el.text:
                    texts.append(t_el.text)
            text = ''.join(texts).strip()
            if text:
                current_paragraphs.append(text)
        elif tag == 'tbl':
            context = '\n'.join(current_paragraphs) if current_paragraphs else ''
            table_contexts[table_idx] = context
            current_paragraphs = []  # Reset for next table
            table_idx += 1

    return table_contexts


def _is_schedule_table(table) -> bool:
    """Determine if a table is a schedule table (has day headers)."""
    if len(table.rows) < 5:
        return False
    first_row = _dedupe_row_cells(table.rows[0])
    text = " ".join(c[0] for c in first_row).lower()
    return any(day.lower() in text for day in DAY_ORDER)


def _count_actual_rooms_per_day(
    table, day_columns: dict[str, tuple[int, int]]
) -> dict[str, int]:
    """Determine the actual number of rooms per day by examining data rows.

    The DOCX grid may allocate more columns for a day than there are real
    rooms (extra columns are used for flexible cell merging).  The true
    room count equals the maximum number of distinct cells observed in any
    single data row for that day.
    """
    max_cells: dict[str, int] = {day: 1 for day in day_columns}

    for row in table.rows[1:]:
        cells = _dedupe_row_cells(row)
        if (
            _is_break_row(cells)
            or _is_footer_row(cells)
            or _is_metadata_row(cells)
        ):
            continue
        if not cells:
            continue
        if _determine_time_block_index(cells[0][0]) is None:
            continue

        for day, (day_col_start, day_col_end) in day_columns.items():
            count = sum(
                1
                for _, cs, ce in cells[1:]
                if cs >= day_col_start and ce <= day_col_end
            )
            if count > max_cells[day]:
                max_cells[day] = count

    return max_cells


def _determine_time_block_index(time_text: str) -> int | None:
    """Match a time label cell to a time block index."""
    # Extract start time from text like "08:30\n~\n10:30\n\n(120 min)"
    match = re.search(r"(\d{1,2}:\d{2})", time_text)
    if not match:
        return None

    start_time = match.group(1)
    # Zero-pad hour
    if len(start_time) == 4:
        start_time = "0" + start_time

    for block in TIME_BLOCKS:
        if block["start"] == start_time:
            return block["index"]
    return None


def parse_docx(
    filepath: str | Path, *, max_tables: int | None = 2,
) -> tuple[list[CellData], list[dict]]:
    """Parse a schedule DOCX file and extract all cell data.

    Args:
        filepath: Path to the DOCX file.
        max_tables: Maximum number of schedule tables to parse.
            The largest tables (by column count) are kept.
            Set to None to parse ALL schedule tables.

    Returns:
        (cells, tables_meta) where:
        - cells: list of CellData for each non-empty schedule cell
        - tables_meta: list of dicts with table metadata (rooms per day)
    """
    doc = Document(str(filepath))

    # ── Context extraction ───────────────────────────────────────
    # Map each table to the paragraph text preceding it.
    table_contexts = _get_table_preceding_paragraphs(doc)

    # ── Room name extraction ────────────────────────────────────
    # Primary: TextBox colour matching (most robust).
    # Fallback: metadata row parsing.
    textbox_rooms = _extract_textbox_rooms(doc)
    online_room_names_fb, offline_room_names_fb = _extract_room_names_from_doc(doc)

    # Find all schedule tables
    schedule_tables = []
    for idx, table in enumerate(doc.tables):
        if _is_schedule_table(table):
            schedule_tables.append((idx, table))

    if not schedule_tables:
        raise ValueError("No schedule tables found in the document")

    # If more tables than limit, take the N largest (by column count)
    if max_tables is not None and len(schedule_tables) > max_tables:
        schedule_tables.sort(key=lambda x: len(x[1].columns), reverse=True)
        schedule_tables = schedule_tables[:max_tables]
        # Re-sort by original index
        schedule_tables.sort(key=lambda x: x[0])

    all_cells = []
    tables_meta = []

    for table_idx, (orig_idx, table) in enumerate(schedule_tables):
        rows = table.rows
        if len(rows) < 2:
            continue

        # Parse header row for day -> column mapping
        header_cells = _dedupe_row_cells(rows[0])
        day_columns = _parse_day_header(header_cells)

        if not day_columns:
            continue

        # Determine actual rooms per day (header span may exceed real rooms)
        actual_rooms = _count_actual_rooms_per_day(table, day_columns)

        day_rooms: dict[str, list[str]] = {}

        # Try colour-based TextBox matching first
        color_matched = _match_rooms_to_table(
            textbox_rooms, table, day_columns, actual_rooms,
        )

        for day, (col_start, col_end) in day_columns.items():
            num_rooms = actual_rooms[day]
            if color_matched and num_rooms <= len(color_matched):
                day_rooms[day] = color_matched[:num_rooms]
            elif table_idx == 0 and online_room_names_fb and num_rooms <= len(online_room_names_fb):
                day_rooms[day] = online_room_names_fb[:num_rooms]
            elif table_idx > 0 and offline_room_names_fb and num_rooms <= len(offline_room_names_fb):
                day_rooms[day] = offline_room_names_fb[:num_rooms]
            else:
                prefix = "Offline" if table_idx > 0 else "Room"
                day_rooms[day] = [
                    f"{prefix} {chr(65 + i)}" for i in range(num_rooms)
                ]

        tables_meta.append(
            {
                "table_index": table_idx,
                "day_columns": day_columns,
                "day_rooms": day_rooms,
                "context_text": table_contexts.get(orig_idx, ''),
            }
        )

        # Parse data rows
        for row_idx in range(1, len(rows)):
            row_cells = _dedupe_row_cells(rows[row_idx])

            if _is_break_row(row_cells) or _is_footer_row(row_cells) or _is_metadata_row(row_cells):
                continue

            # First cell should be the time label
            if not row_cells:
                continue

            time_cell_text = row_cells[0][0]
            tb_index = _determine_time_block_index(time_cell_text)

            if tb_index is None:
                continue

            time_block = TIME_BLOCKS[tb_index]

            # Process cells grouped by day with ordinal room mapping
            for day, (day_col_start, day_col_end) in day_columns.items():
                grid_cols = day_col_end - day_col_start
                num_actual = actual_rooms[day]

                # Collect cells for this day, sorted by column position
                day_data_cells = sorted(
                    [
                        (t, cs, ce)
                        for t, cs, ce in row_cells[1:]
                        if cs >= day_col_start and ce <= day_col_end
                    ],
                    key=lambda x: x[1],
                )

                if not day_data_cells:
                    continue

                running_room = 0
                for text, col_start, col_end in day_data_cells:
                    cell_span = col_end - col_start

                    # Determine room indices for this cell
                    if cell_span >= grid_cols:
                        # Cell spans all grid columns → all rooms
                        room_indices = list(range(num_actual))
                        running_room = num_actual
                    elif len(day_data_cells) == num_actual:
                        # Exactly one cell per room → sequential
                        room_indices = [running_room]
                        running_room += 1
                    else:
                        # Proportional assignment based on grid span
                        rooms_count = max(
                            1, round(cell_span * num_actual / grid_cols)
                        )
                        rooms_count = min(
                            rooms_count, max(1, num_actual - running_room)
                        )
                        room_indices = list(
                            range(running_room, running_room + rooms_count)
                        )
                        running_room += rooms_count

                    if not text.strip():
                        continue

                    cell_data = CellData(
                        text=text,
                        day=day,
                        room_indices=room_indices,
                        time_block_index=tb_index,
                        time_block_start=time_block["start"],
                        time_block_end=time_block["end"],
                        time_block_duration=time_block["duration"],
                        table_index=table_idx,
                    )
                    all_cells.append(cell_data)

    return all_cells, tables_meta


def build_room_list(tables_meta: list[dict]) -> dict[str, list[RoomInfo]]:
    """Build a unified room list per day from all tables.

    Returns dict: day_name -> list of RoomInfo (in order for grid columns).
    """
    day_rooms: dict[str, list[RoomInfo]] = {}

    for meta in tables_meta:
        for day, rooms in meta["day_rooms"].items():
            if day not in day_rooms:
                day_rooms[day] = []
            for ri, room_name in enumerate(rooms):
                day_rooms[day].append(
                    RoomInfo(
                        name=room_name,
                        table_index=meta["table_index"],
                        room_index_in_table=ri,
                    )
                )

    return day_rooms


def compute_room_global_col(
    cell: CellData,
    day_rooms: list[RoomInfo],
) -> tuple[int, int]:
    """Compute global grid column range for a cell's rooms.

    Returns (col_start, col_end) as 1-indexed grid columns
    (col 1 = time label, col 2 = first room).
    """
    matching_indices = []
    for global_idx, room_info in enumerate(day_rooms):
        if room_info.table_index == cell.table_index and room_info.room_index_in_table in cell.room_indices:
            matching_indices.append(global_idx)

    if not matching_indices:
        return (2, 3)  # fallback: first room

    col_start = min(matching_indices) + 2  # +2: col 1 = time label
    col_end = max(matching_indices) + 3  # exclusive
    return (col_start, col_end)
